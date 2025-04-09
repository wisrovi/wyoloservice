import os
import cv2
import matplotlib.pyplot as plt
import hashlib
from ultralytics import YOLO
from PIL import Image
import numpy as np
from datetime import datetime
import seaborn as sns
from loguru import logger
import random
import yaml
import argparse
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from docx2pdf import convert

random.seed(33)

CLASSIFICATION = "Clasificación"
DETECTION = "Detección"
SEGMENTATION = "Segmentación"

CLASS_IMBALANCE_RATIO = 1.5  # Max ratio between most and least common class
IMAGE_SIZE_VARIATION_THRESHOLD = 400  # px difference in width or height
BBOX_AREA_VARIATION_RATIO = 20  # Ratio of largest to smallest bounding box area
ASPECT_RATIO_DIVERSITY_THRESHOLD = 50  # Number of unique aspect ratios (rounded)
BBOX_ASPECT_RATIO_VARIANCE = 2.0  # Difference between max and min aspect ratio
CENTER_POSITION_BIN_COUNT = 75 # Min number of unique rounded x or y positions
BBOX_WH_CORRELATION_THRESHOLD = 0.3  # Min correlation to consider width/height related

class YOLODataValidator:
    def __init__(self, image_folder, label_folder, data_yaml, dataset_type, dataset_path, project_name, author_name, content_description, verbose=False):
        self.image_folder = image_folder
        self.label_folder = label_folder
        self.data_yaml = data_yaml
        self.dataset_type = dataset_type
        self.project_name = project_name
        self.author_name = author_name
        self.content_description = content_description
        self.report_folder = f"{dataset_path}/reports"
        self.word_path = f"{self.report_folder}/EDA Report.docx"
        self.model_clas = "yolov8n-cls.pt"
        self.model_seg = "yolov8m-seg.pt"
        self.model_det = "yolov10x.pt"
        self.verbose = verbose
        os.makedirs(self.report_folder, exist_ok=True)
        self._configure_logger()

    def _configure_logger(self):
        logger.remove()
        log_level = "DEBUG" if self.verbose else "INFO"
        logger.add(
            os.path.join(self.report_folder, "validation.log"),
            level=log_level,
            rotation="1 day",
            retention="7 days",
        )
        logger.add(lambda msg: print(msg, end=""), level=log_level)

    def analyze_class_distribution(self):
        class_counts = {}

        if self.dataset_type == CLASSIFICATION:
            class_names = sorted(os.listdir(self.image_folder))
            for class_name in class_names:
                class_path = os.path.join(self.image_folder, class_name)
                if os.path.isdir(class_path):
                    num_images = len([f for f in os.listdir(class_path) if f.endswith((".jpg", ".png", ".jpeg"))])
                    class_counts[class_name] = num_images
        else:
            with open(self.data_yaml, 'r') as f:
                data = yaml.safe_load(f)
            
            class_names = data['names']

            for label_file in os.listdir(self.label_folder):
                if label_file.endswith(".txt"):
                    with open(os.path.join(self.label_folder, label_file), "r") as f:
                        for line in f:
                            class_id = int(float(line.split()[0]))
                            class_name = class_names[class_id]
                            class_counts[class_name] = class_counts.get(class_name, 0) + 1

        sorted_class_names = sorted(class_counts.keys(), key=lambda x: class_names.index(x) if x in class_names else -1)

        # Graficar la distribución de clases
        plt.figure(figsize=(12, 12))
        sns.barplot(x=sorted_class_names, y=[class_counts[name] for name in sorted_class_names])
        plt.xlabel("Class")
        plt.ylabel("Count")
        plt.title("Class Distribution")
        plt.xticks(rotation=45)
        plt.savefig(os.path.join(self.report_folder, "class_distribution.png"))
        plt.close()

        return class_counts

    def analyze_image_sizes(self):
        widths, heights = [], []

        if self.dataset_type == CLASSIFICATION:
            for class_folder in os.listdir(self.image_folder):
                class_path = os.path.join(self.image_folder, class_folder)
                
                if os.path.isdir(class_path):
                    for image_file in os.listdir(class_path):
                        if image_file.endswith((".jpg", ".png", ".jpeg")):
                            with Image.open(os.path.join(class_path, image_file)) as img:
                                widths.append(img.width)
                                heights.append(img.height)
        else:
            for image_file in os.listdir(self.image_folder):
                if image_file.endswith((".jpg", ".png", ".jpeg")):
                    with Image.open(os.path.join(self.image_folder, image_file)) as img:
                        widths.append(img.width)
                        heights.append(img.height)

        plt.figure(figsize=(10, 6))
        sns.histplot(widths, bins=30, color="blue", label="Width")
        sns.histplot(heights, bins=30, color="red", label="Height")
        plt.xlabel("Pixels")
        plt.ylabel("Frequency")
        plt.title("Image Size Distribution")
        plt.legend()
        plt.savefig(os.path.join(self.report_folder, "image_size_distribution.png"))
        plt.close()

        return widths, heights

    def analyze_bbox_areas(self):
        areas = []

        if self.dataset_type == CLASSIFICATION:
            for class_folder in os.listdir(self.image_folder):
                class_path = os.path.join(self.image_folder, class_folder)

                if os.path.isdir(class_path):
                    for image_file in os.listdir(class_path):
                        if image_file.endswith((".jpg", ".png", ".jpeg")):
                            with Image.open(os.path.join(class_path, image_file)) as img:
                                width, height = img.size
                                areas.append(width * height)
        else:
            for label_file in os.listdir(self.label_folder):
                if label_file.endswith(".txt"):
                    with open(os.path.join(self.label_folder, label_file), "r") as f:
                        for line in f:
                            values = list(map(float, line.split()))
                            # Si hay más de 5 valores, puedes suponer que hay más de un objeto o información extra
                            num_values = len(values)
                            if num_values >= 5:
                                # Extraer los 5 primeros valores esperados para cada objeto
                                _, x_center, y_center, box_width, box_height = values[:5]
                                areas.append(box_width * box_height)
                            else:
                                logger.warning(f"Línea con datos insuficientes en {label_file}: {line}")
        
        plt.figure(figsize=(10, 6))
        sns.histplot(areas, bins=30, color="green")
        plt.xlabel("Bounding Box Area (normalized)")
        plt.ylabel("Frequency")
        plt.title("Bounding Box Area Distribution")
        plt.savefig(os.path.join(self.report_folder, "bbox_area_distribution.png"))
        plt.close()

        return areas

    def analyze_aspect_ratios(self):
        aspect_ratios = []

        if self.dataset_type == CLASSIFICATION:
            for class_folder in os.listdir(self.image_folder):
                class_path = os.path.join(self.image_folder, class_folder)

                if os.path.isdir(class_path):
                    for image_file in os.listdir(class_path):
                        if image_file.endswith((".jpg", ".png", ".jpeg")):
                            image_path = os.path.join(class_path, image_file)
                            try:
                                with Image.open(image_path) as img:
                                    aspect_ratio = img.width / img.height
                                    aspect_ratios.append(aspect_ratio)
                            except Exception as e:
                                logger.exception(f"Error al abrir la imagen {image_file}: {e}")

        else:
            for image_file in os.listdir(self.image_folder):
                if image_file.endswith((".jpg", ".png", ".jpeg")):
                    with Image.open(os.path.join(self.image_folder, image_file)) as img:
                        aspect_ratios.append(img.width / img.height)

        plt.figure(figsize=(10, 6))
        sns.histplot(aspect_ratios, bins=30, color="purple")
        plt.xlabel("Aspect Ratio (Width/Height)")
        plt.ylabel("Frequency")
        plt.title("Image Aspect Ratio Distribution")
        plt.savefig(os.path.join(self.report_folder, "aspect_ratio_distribution.png"))
        plt.close()

        return aspect_ratios

    def get_all_images(self, folder, split_name):
            return [
                (os.path.join(root, f), split_name)
                for root, _, files in os.walk(folder)
                for f in files
                if f.lower().endswith((".jpg", ".jpeg", ".png"))
            ]

    def detect_duplicates_and_overlaps(self):
        image_hashes = {}
        duplicates = []
        image_paths = []

        if self.dataset_type == CLASSIFICATION:
            for split_folder in os.listdir(dataset_path):
                if split_folder in ['train', 'test', 'val']:
                    split_path = os.path.join(dataset_path, split_folder)
                    if os.path.isdir(split_path):
                        image_paths.extend(self.get_all_images(split_path, split_folder))
        else:
            with open(self.data_yaml, 'r') as f:
                data = yaml.safe_load(f)
            for split in ['train', 'val', 'test']:
                split_path = data.get(split)
                if split_path and os.path.isdir(split_path):
                    image_paths.extend(self.get_all_images(split_path, split))

        for img_path, split in image_paths:
            img = cv2.imread(img_path)
            if img is not None:
                img_hash = hashlib.sha1(img).hexdigest()
                if img_hash in image_hashes:
                    duplicates.append(((os.path.basename(img_path), split), 
                                        (os.path.basename(image_hashes[img_hash][0]), image_hashes[img_hash][1])))
                else:
                    image_hashes[img_hash] = (img_path, split)

        return duplicates

    def validate_yolo_format(self, data_yaml):
        model_name = ""

        if self.dataset_type == CLASSIFICATION:
            model_name = self.model_clas
        if self.dataset_type == SEGMENTATION:
            model_name = self.model_seg
        if self.dataset_type == DETECTION:
            model_name = self.model_det
            
        model = YOLO(model_name)
        
        try:
            results = model.val(
                data=data_yaml,
                save=True,
                verbose=False,
                plots=True,
                project=os.path.join(self.report_folder, "runs"),
                name="val",
            )
            return results.results_dict
        except Exception as e:
            logger.exception(f"Error durante la validación: {e}")
        finally:
            del model
            try:
                if os.path.exists(model_name):
                    os.remove(model_name)
            except Exception as e:
                logger.exception(f"Error eliminando el modelo temporal: {e}")
        
    def validate_image_quality(self):
        corrupt_images, small_images = [], []

        if self.dataset_type == CLASSIFICATION:
                for class_folder in os.listdir(self.image_folder):
                    class_path = os.path.join(self.image_folder, class_folder)
                    
                    if os.path.isdir(class_path):
                        for image_file in os.listdir(class_path):
                            if image_file.endswith((".jpg", ".png", ".jpeg")):
                                image_path = os.path.join(class_path, image_file)
                                try:
                                    img = cv2.imread(image_path)
                                    if img is None:
                                        corrupt_images.append(image_file)
                                    else:
                                        height, width, _ = img.shape
                                        if width < 64 or height < 64:
                                            small_images.append((image_file, width, height))
                                except Exception:
                                    corrupt_images.append(image_file)
        else:
            for image_file in os.listdir(self.image_folder):
                if image_file.endswith((".jpg", ".png", ".jpeg")):
                    image_path = os.path.join(self.image_folder, image_file)
                    try:
                        img = cv2.imread(image_path)
                        if img is None:
                            corrupt_images.append(image_file)
                        else:
                            height, width, _ = img.shape
                            if width < 64 or height < 64:
                                small_images.append((image_file, width, height))
                    except Exception:
                        corrupt_images.append(image_file)

        return corrupt_images, small_images

    def generate_example_mosaics(self, num_mosaics=3, mosaic_size=(3, 3), target_size=(224, 224)):
        image_files = []
        class_names = []
        class_colors = {}
        mosaics = []

        if dataset_type == CLASSIFICATION:
            for subdir, _, files in os.walk(self.image_folder):
                for file in files:
                    if file.endswith((".jpg", ".png")):
                        image_files.append(os.path.join(subdir, file))

                        class_name = os.path.basename(subdir)
                        if class_name not in class_names:
                            class_names.append(class_name)

            for class_name in class_names:
                class_colors[class_name] = tuple(random.randint(0, 255) for _ in range(3))

            for mosaic_idx in range(num_mosaics):
                mosaic_images = []
                random.shuffle(image_files)

                for img_file in random.sample(image_files, min(mosaic_size[0] * mosaic_size[1], len(image_files))):
                    img_path = img_file
                    img = cv2.imread(img_path)
                    if img is None:
                        continue

                    # Redimensionar la imagen al tamaño objetivo
                    img_resized = cv2.resize(img, target_size)

                    if dataset_type == CLASSIFICATION:
                        class_name = os.path.basename(os.path.dirname(img_file))
                        class_color = class_colors.get(class_name, (255, 255, 255))

                        text_position = (10, 30)
                        font_scale = 0.8
                        thickness = 2
                        font = cv2.FONT_HERSHEY_SIMPLEX

                        (text_width, text_height), _ = cv2.getTextSize(class_name, font, font_scale, thickness)

                        cv2.rectangle(
                            img_resized,
                            (text_position[0] - 5, text_position[1] - text_height - 5),
                            (text_position[0] + text_width + 5, text_position[1] + 5),
                            class_color, -1
                        )

                        cv2.putText(img_resized, class_name, text_position, font, font_scale, (255, 255, 255), thickness)

                    mosaic_images.append(img_resized)

                if mosaic_images:
                    height, width, _ = mosaic_images[0].shape
                    mosaic_height = height * mosaic_size[0]
                    mosaic_width = width * mosaic_size[1]
                    mosaic = np.zeros((mosaic_height, mosaic_width, 3), dtype=np.uint8)

                    row_offset = 0
                    for row_idx in range(mosaic_size[0]):
                        col_offset = 0
                        for col_idx in range(mosaic_size[1]):
                            img_idx = row_idx * mosaic_size[1] + col_idx
                            if img_idx >= len(mosaic_images):
                                break
                            img = mosaic_images[img_idx]
                            mosaic[row_offset: row_offset + img.shape[0], col_offset: col_offset + img.shape[1]] = img
                            col_offset += img.shape[1]
                        row_offset += mosaic_images[row_idx].shape[0] if row_idx < len(mosaic_images) else 0

                    mosaic_filename = os.path.join(self.report_folder, f"mosaic_{mosaic_idx}.png")
                    cv2.imwrite(mosaic_filename, mosaic)
                    mosaics.append(mosaic_filename)

        else:  # DETECTION or SEGMENTATION
            image_files = [f for f in os.listdir(self.image_folder) if f.endswith((".jpg", ".png", ".jpeg"))]
            label_files = {
                f.replace(".jpg", ".txt").replace(".png", ".txt").replace(".jpeg", ".txt"): f
                for f in os.listdir(self.label_folder) if f.endswith(".txt")
            }

            for mosaic_idx in range(num_mosaics):
                mosaic_images = []
                random.shuffle(image_files)

                for img_file in random.sample(image_files, min(mosaic_size[0] * mosaic_size[1], len(image_files))):
                    img_path = os.path.join(self.image_folder, img_file)
                    img = cv2.imread(img_path)
                    if img is None:
                        continue
                    img_resized = cv2.resize(img, target_size)

                    if self.dataset_type == "CLASSIFICATION":
                        class_name = os.path.basename(os.path.dirname(img_file))
                        class_color = class_colors.get(class_name, (255, 255, 255))
                        text_position = (10, 30)
                        cv2.rectangle(img_resized, (5, 10), (150, 40), class_color, -1)
                        cv2.putText(img_resized, class_name, text_position, cv2.FONT_HERSHEY_SIMPLEX, 0.8, (255, 255, 255), 1, lineType=cv2.LINE_AA)
                    else:
                        label_file = img_file.replace(".jpg", ".txt").replace(".png", ".txt").replace(".jpeg", ".txt")
                        if label_file not in label_files:
                            continue
                        label_path = os.path.join(self.label_folder, label_file)
                        labels = []
                        with open(label_path, "r") as f:
                            for line in f:
                                values = line.split()
                                if len(values) == 5:  # Detection: class_id, x_center, y_center, width, height
                                    class_id = int(float(values[0]))
                                    x_center, y_center, w, h = map(float, values[1:])
                                    img_h, img_w, _ = img_resized.shape
                                    x_min = int((x_center - w / 2) * img_w)
                                    y_min = int((y_center - h / 2) * img_h)
                                    x_max = int((x_center + w / 2) * img_w)
                                    y_max = int((y_center + h / 2) * img_h)
                                    labels.append((class_id, (x_min, y_min, x_max, y_max)))
                                elif len(values) > 5:  # Segmentation: class_id + polygon points
                                    class_id = int(float(values[0]))
                                    polygon_coords = list(map(float, values[1:]))
                                    polygon_points = [(polygon_coords[i], polygon_coords[i + 1]) for i in range(0, len(polygon_coords), 2)]
                                    labels.append((class_id, polygon_points))

                        for label in labels:
                            if isinstance(label[1], tuple):  # Bounding Box (Detection)
                                class_id, (x_min, y_min, x_max, y_max) = label
                                cv2.rectangle(img_resized, (x_min, y_min), (x_max, y_max), (0, 255, 0), 1)
                                cv2.putText(img_resized, str(class_id), (x_min, y_min - 5), cv2.FONT_HERSHEY_SIMPLEX, 0.3, (0, 0, 255), 1, lineType=cv2.LINE_AA)
                            else:  # Polygon (Segmentation)
                                class_id, polygon_points = label
                                polygon_points = np.array(polygon_points, dtype=np.float32)
                                polygon_points[:, 0] *= img_resized.shape[1]
                                polygon_points[:, 1] *= img_resized.shape[0]
                                polygon_points = polygon_points.astype(np.int32)
                                cv2.polylines(img_resized, [polygon_points], isClosed=True, color=(0, 255, 0), thickness=1)
                                text_position = tuple(polygon_points[0])
                                # Verificar que el texto no esté fuera de los límites de la imagen
                                if text_position[1] - 10 < 0:
                                    text_position = (text_position[0], text_position[1] + 20)
                                if text_position[0] + 40 > img_resized.shape[1]:
                                    text_position = (text_position[0] - 40, text_position[1])
                                cv2.putText(img_resized, str(class_id), text_position, cv2.FONT_HERSHEY_SIMPLEX, 0.5, (0, 0, 255), 1, lineType=cv2.LINE_AA)
                            
                        # Asegurar que el texto no se salga de los límites de la imagen
                        image_name_position = (10, 30)
                        if image_name_position[1] + 10 > img_resized.shape[0]:
                            image_name_position = (10, img_resized.shape[0] - 10)

                        cv2.putText(
                            img_resized,
                            img_file,
                            image_name_position,
                            cv2.FONT_HERSHEY_SIMPLEX,
                            0.3,
                            (0, 255, 255),  # Amarillo
                            1,
                            lineType=cv2.LINE_AA  # Suaviza las líneas
                        )

                    mosaic_images.append(img_resized)

                if mosaic_images:
                    height, width, _ = mosaic_images[0].shape
                    mosaic = np.zeros((height * mosaic_size[0], width * mosaic_size[1], 3), dtype=np.uint8)
                    for row in range(mosaic_size[0]):
                        for col in range(mosaic_size[1]):
                            img_idx = row * mosaic_size[1] + col
                            if img_idx < len(mosaic_images):
                                img = mosaic_images[img_idx]
                                mosaic[row * height:(row + 1) * height, col * width:(col + 1) * width] = img

                    mosaic_filename = os.path.join(self.report_folder, f"mosaic_{mosaic_idx}.png")
                    cv2.imwrite(mosaic_filename, mosaic)
                    mosaics.append(mosaic_filename)

        return mosaics


    def analyze_bbox_aspect_ratios(self):
        aspect_ratios = []

        if self.dataset_type == CLASSIFICATION:
            for class_folder in os.listdir(self.image_folder):
                class_path = os.path.join(self.image_folder, class_folder)

                if os.path.isdir(class_path):
                    for image_file in os.listdir(class_path):
                        if image_file.endswith((".jpg", ".png", ".jpeg")):
                            image_path = os.path.join(class_path, image_file)
                            with Image.open(image_path) as img:
                                width, height = img.size
                                if height != 0:
                                    aspect_ratios.append(width / height)
        else:
            for label_file in os.listdir(self.label_folder):
                if label_file.endswith(".txt"):
                    with open(os.path.join(self.label_folder, label_file), "r") as f:
                        for line in f:
                            values = list(map(float, line.split()))
                            
                            if len(values) >= 5:
                                _, _, _, box_width, box_height = values[:5]
                                
                                if box_height != 0:
                                    aspect_ratios.append(box_width / box_height)
        
        plt.figure(figsize=(10, 6))
        sns.histplot(aspect_ratios, bins=30, color="orange")
        plt.xlabel("Bounding Box Aspect Ratio (Width/Height)")
        plt.ylabel("Frequency")
        plt.title("Bounding Box Aspect Ratio Distribution")
        plt.savefig(os.path.join(self.report_folder, "bbox_aspect_ratio_distribution.png"))
        plt.close()
        
        return aspect_ratios


    def analyze_bbox_center_positions(self):
        x_centers, y_centers = [], []

        if self.dataset_type == CLASSIFICATION:
            for class_folder in os.listdir(self.image_folder):
                class_path = os.path.join(self.image_folder, class_folder)
                
                if os.path.isdir(class_path):
                    for image_file in os.listdir(class_path):
                        if image_file.endswith((".jpg", ".png", ".jpeg")):
                            image_path = os.path.join(class_path, image_file)
                            with Image.open(image_path) as img:
                                x_centers.append(0.5)  # El centro de la imagen es siempre 0.5 para clasificación
                                y_centers.append(0.5)  # El centro de la imagen es siempre 0.5 para clasificación
        else:
            for label_file in os.listdir(self.label_folder):
                if label_file.endswith(".txt"):
                    with open(os.path.join(self.label_folder, label_file), "r") as f:
                        for line in f:
                            values = list(map(float, line.split()))
                            num_values = len(values)
                            if num_values >= 5:
                                _, x_center, y_center, _, _ = values[:5]
                                x_centers.append(x_center)
                                y_centers.append(y_center)
                            else:
                                logger.warning(f"Línea con datos insuficientes en {label_file}: {line}")

        plt.figure(figsize=(12, 6))
        plt.subplot(1, 2, 1)
        sns.histplot(x_centers, bins=30, color="purple")
        plt.xlabel("Bounding Box X Center (normalized)")
        plt.ylabel("Frequency")
        plt.title("Bounding Box X Center Distribution")
        plt.subplot(1, 2, 2)
        sns.histplot(y_centers, bins=30, color="blue")
        plt.xlabel("Bounding Box Y Center (normalized)")
        plt.ylabel("Frequency")
        plt.title("Bounding Box Y Center Distribution")
        plt.savefig(
            os.path.join(self.report_folder, "bbox_center_position_distribution.png")
        )
        plt.close()

        return x_centers, y_centers

    def analyze_bbox_width_height(self):
        widths, heights = [], []

        if self.dataset_type == CLASSIFICATION:
            for class_folder in os.listdir(self.image_folder):
                class_path = os.path.join(self.image_folder, class_folder)

                if os.path.isdir(class_path):
                    for image_file in os.listdir(class_path):
                        if image_file.endswith((".jpg", ".png", ".jpeg")):
                            image_path = os.path.join(class_path, image_file)
                            with Image.open(image_path) as img:
                                widths.append(img.width)
                                heights.append(img.height)
        else:
            for label_file in os.listdir(self.label_folder):
                if label_file.endswith(".txt"):
                    with open(os.path.join(self.label_folder, label_file), "r") as f:
                        for line in f:
                            values = list(map(float, line.split()))
                            num_values = len(values)
                            if num_values >= 5:
                                _, _, _, box_width, box_height = values[:5]
                                widths.append(box_width)
                                heights.append(box_height)
                            else:
                                logger.warning(f"Línea con datos insuficientes en {label_file}: {line}")

        plt.figure(figsize=(8, 8))
        plt.scatter(widths, heights, alpha=0.5)
        plt.xlabel("Bounding Box Width (normalized)")
        plt.ylabel("Bounding Box Height (normalized)")
        plt.title("Bounding Box Width vs Height")
        plt.grid(True)
        plt.savefig(os.path.join(self.report_folder, "bbox_width_height.png"))
        plt.close()

        return widths, heights

    # Índice Word
    def add_index(self, doc, section_titles):
        # Crear el índice en una lista temporal de elementos
        index_doc = Document()
        index_doc.add_page_break()
        index_doc.add_paragraph("Índice", style="Heading 1")
        for i, title in enumerate(section_titles, 1):
            index_doc.add_paragraph(f"{i}. {title}", style="Normal")
        index_doc.add_page_break()

        index_elements = [p._element for p in index_doc.paragraphs]
        index_elements.append(index_doc.paragraphs[-1]._element.getparent())

        # Insertar después del título, modificar si se añaden datos antes del índice
        insert_pos = 8
        body = doc._body._element
        for element in reversed(index_elements):
            body.insert(insert_pos, element)

    # Generar el reporte en Word
    def create_word_report(self,
        class_distribution, image_dimensions, bbox_areas, aspect_ratios, duplicates,
        validation_results, corrupt_images, small_images, bbox_aspect_ratios, 
        bbox_center_positions, bbox_width_height):
        
        try:
            doc = Document()
            section_titles = []
            
            # Portada
            doc.add_paragraph(f"Análisis Exploratorio de Datos (EDA) {self.dataset_type}", style="Title").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph("Informe de Análisis Exploratorio de Datos", style="Heading 1").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph("\n", style="Normal")
            doc.add_paragraph(f"Autor: {self.author_name}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            doc.add_paragraph(f"Proyecto: {self.project_name}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            doc.add_paragraph(f"Modelo: {self.dataset_type}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            doc.add_paragraph(f"Descripción: {self.content_description}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}", style="Normal").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            doc.add_paragraph("\n", style="Normal")

            # Secciones del informe
            doc.add_paragraph("Distribución de Clases:", style="Heading 1")
            section_titles.append("Distribución de Clases")
            doc.add_picture(os.path.join(self.report_folder, "class_distribution.png"), width=Inches(6))

            doc.add_page_break()
            doc.add_paragraph("Distribución de Tamaños de Imágenes:", style="Heading 1")
            section_titles.append("Distribución de Tamaños de Imágenes")
            doc.add_picture(os.path.join(self.report_folder, "image_size_distribution.png"), width=Inches(6))

            doc.add_page_break()
            doc.add_paragraph("Distribución de Áreas de Bounding Boxes:", style="Heading 1")
            section_titles.append("Distribución de Áreas de Bounding Boxes")
            doc.add_picture(os.path.join(self.report_folder, "bbox_area_distribution.png"), width=Inches(6))

            doc.add_page_break()
            doc.add_paragraph("Distribución de Relaciones de Aspecto:", style="Heading 1")
            section_titles.append("Distribución de Relaciones de Aspecto")
            doc.add_picture(os.path.join(self.report_folder, "aspect_ratio_distribution.png"), width=Inches(6))

            doc.add_page_break()
            doc.add_paragraph("Distribución de Relaciones de Aspecto de Bounding Boxes:", style="Heading 1")
            section_titles.append("Distribución de Relaciones de Aspecto de Bounding Boxes")
            doc.add_picture(os.path.join(self.report_folder, "bbox_aspect_ratio_distribution.png"), width=Inches(6))

            doc.add_page_break()
            doc.add_paragraph("Distribución de Posiciones Centrales de Bounding Boxes:", style="Heading 1")
            section_titles.append("Distribución de Posiciones Centrales de Bounding Boxes")
            doc.add_picture(os.path.join(self.report_folder, "bbox_center_position_distribution.png"), width=Inches(6))

            doc.add_page_break()
            doc.add_paragraph("Gráfico de Dispersión de Ancho vs. Alto de Bounding Boxes:", style="Heading 1")
            section_titles.append("Gráfico de Dispersión de Ancho vs. Alto de Bounding Boxes")
            doc.add_picture(os.path.join(self.report_folder, "bbox_width_height.png"), width=Inches(6))

            doc.add_page_break()
            doc.add_paragraph("Mosaicos de Imágenes de Ejemplo:", style="Heading 1")
            section_titles.append("Mosaicos de Imágenes de Ejemplo")
            mosaics = self.generate_example_mosaics()
            for mosaic_file in mosaics:
                doc.add_picture(mosaic_file, width=Inches(6))
                doc.add_paragraph("")

            doc.add_page_break()
            
            doc.add_paragraph("Resultados de Validación con Ultralytics:", style="Heading 1")
            section_titles.append("Resultados de Validación con Ultralytics")
            for key, value in validation_results.items():
                doc.add_paragraph(f"{key}: {round(float(value), 4)}", style="Normal")
            doc.add_page_break()

            doc.add_paragraph("Validación de Calidad de Imágenes:", style="Heading 1")
            section_titles.append("Validación de Calidad de Imágenes")
            if corrupt_images:
                doc.add_paragraph("Imágenes corruptas encontradas:", style="Heading 2")
                for img in corrupt_images:
                    doc.add_paragraph(f"- {img}")
            else:
                doc.add_paragraph("No se encontraron imágenes corruptas.")

            if small_images:
                doc.add_paragraph("Imágenes con dimensiones pequeñas encontradas:", style="Heading 2")
                for img, width, height in small_images:
                    doc.add_paragraph(f"- {img} ({width}x{height})")
            else:
                doc.add_paragraph("No se encontraron imágenes con dimensiones pequeñas.")

            if duplicates:
                doc.add_paragraph("Imágenes Duplicadas:", style="Heading 2")
                for dup in duplicates:
                    img1, split1 = dup[0]
                    img2, split2 = dup[1]
                    doc.add_paragraph(f"- '{img1}' en '{split1}' y '{img2}' en '{split2}'")
            else:
                doc.add_paragraph("No se encontraron imágenes duplicadas.")

            doc.add_page_break()

            # Sección de Conclusiones
            doc.add_paragraph("Conclusiones del Análisis EDA:", style="Heading 1")
            section_titles.append("Conclusiones del Análisis EDA")

            conclusion_lines = []

            # Imágenes corruptas
            if corrupt_images:
                conclusion_lines.append("Se detectaron imágenes corruptas que deben ser eliminadas o reemplazadas.")
            else:
                conclusion_lines.append("No se encontraron imágenes corruptas.")

            # Imágenes pequeñas
            if small_images:
                conclusion_lines.append("Existen imágenes con dimensiones pequeñas que podrían afectar el rendimiento del modelo.")
            else:
                conclusion_lines.append("Todas las imágenes tienen dimensiones adecuadas.")

            # Imágenes duplicadas
            if duplicates:
                conclusion_lines.append("Se encontraron imágenes duplicadas, lo que requiere una revisión y limpieza de estos casos.")
            else:
                conclusion_lines.append("No se detectaron imágenes duplicadas.")

            # Distribución de clases
            if class_distribution:
                class_counts = list(class_distribution.values())
                max_count = max(class_counts)
                min_count = min(class_counts)
                balance_ratio = max_count / (min_count + 1e-5)

                max_class = [name for name, count in class_distribution.items() if count == max_count][0]
                min_class = [name for name, count in class_distribution.items() if count == min_count][0]

                if balance_ratio > CLASS_IMBALANCE_RATIO:
                    conclusion_lines.append(f"La distribución de clases está desbalanceada. La clase '{max_class}' tiene {max_count} etiquetas, "
                                            f"mientras que la clase '{min_class}' tiene {min_count} etiquetas. (Gráfica: Distribución de Clases)")
                else:
                    conclusion_lines.append("La distribución de clases es equilibrada. (Gráfica: Distribución de Clases)")

            # Tamaños de imágenes
            if image_dimensions:
                widths = [dim[0] for dim in image_dimensions]
                heights = [dim[1] for dim in image_dimensions]
                # Calcular la diferencia máxima de tamaño
                max_width_diff = max(widths) - min(widths)
                max_height_diff = max(heights) - min(heights)
                if max(widths) - min(widths) > IMAGE_SIZE_VARIATION_THRESHOLD or max(heights) - min(heights) > IMAGE_SIZE_VARIATION_THRESHOLD:
                    conclusion_lines.append(f"Se observa variabilidad significativa en los tamaños de las imágenes. "
                                            f"La diferencia máxima en el ancho es de {max_width_diff} píxeles, "
                                            f"y la diferencia máxima en la altura es de {max_height_diff} píxeles. "
                                            f"(Gráfica: Distribución de Tamaños de Imágenes)")
                else:
                    conclusion_lines.append("Las dimensiones de las imágenes son similares. (Gráfica: Distribución de Tamaños de Imágenes)")

            # Bounding box area
            if bbox_areas:
                if max(bbox_areas) / (min(bbox_areas) + 1e-5) > BBOX_AREA_VARIATION_RATIO:
                    conclusion_lines.append("Se observa variabilidad en las áreas de los bounding boxes, lo que sugiere la presencia de objetos de diferentes tamaños. "
                                            "(Gráfica: Distribución de Áreas de Bounding Boxes)")
                else:
                    conclusion_lines.append("Las áreas de los bounding boxes son relativamente uniformes. (Gráfica: Distribución de Áreas de Bounding Boxes)")

            # Aspect ratios
            if aspect_ratios:
                unique_ratios = len(set(round(r, 1) for r in aspect_ratios))
                if unique_ratios > ASPECT_RATIO_DIVERSITY_THRESHOLD:
                    conclusion_lines.append("Hay una alta diversidad en las relaciones de aspecto, lo cual puede exigir un mayor nivel de flexibilidad en el modelo. "
                                            "(Gráfica: Distribución de Relaciones de Aspecto)")
                else:
                    conclusion_lines.append("Las relaciones de aspecto son consistentes, lo que puede facilitar el entrenamiento del modelo. "
                                            "(Gráfica: Distribución de Relaciones de Aspecto)")

            # Bounding box aspect ratios
            if bbox_aspect_ratios:
                ratio_variance = max(bbox_aspect_ratios) - min(bbox_aspect_ratios)
                if ratio_variance > BBOX_ASPECT_RATIO_VARIANCE:
                    conclusion_lines.append("Las relaciones de aspecto de los bounding boxes varían significativamente, lo que indica una diversidad en la forma "
                                            "de los objetos detectados en las imágenes. (Gráfica: Distribución de Relaciones de Aspecto de Bounding Boxes)")
                else:
                    conclusion_lines.append("Las relaciones de aspecto de los bounding boxes son consistentes. "
                                            "(Gráfica: Distribución de Relaciones de Aspecto de Bounding Boxes)")

            # Bounding box center positions
            if bbox_center_positions:
                xs = [p[0] for p in bbox_center_positions]
                ys = [p[1] for p in bbox_center_positions]
                if len(set(round(x, 1) for x in xs)) < CENTER_POSITION_BIN_COUNT or len(set(round(y, 1) for y in ys)) < CENTER_POSITION_BIN_COUNT:
                    conclusion_lines.append(
                                "Las posiciones centrales de las etiquetas se agrupan en áreas específicas. Este patrón de agrupación no necesariamente refleja un sesgo, "
                                "sino que podría una característica propia del dataset. "
                                "(Gráfica: Distribución de Posiciones Centrales de Bounding Boxes)"
                            )
                else:
                    conclusion_lines.append("Las posiciones centrales de las etiquetas están distribuidas en las imágenes. "
                                            "(Gráfica: Distribución de Posiciones Centrales de Bounding Boxes)")

            # Bounding box width vs height
            if bbox_width_height:
                ws = [wh[0] for wh in bbox_width_height]
                hs = [wh[1] for wh in bbox_width_height]
                if len(ws) > 0 and len(hs) > 0:
                    corr = np.corrcoef(ws, hs)[0, 1]
                    if corr < BBOX_WH_CORRELATION_THRESHOLD:
                        conclusion_lines.append("Existe una baja correlación entre el ancho y el alto de los bounding boxes. "
                                                "(Gráfica: Dispersión de Ancho vs. Alto de Bounding Boxes)")
                    else:
                        conclusion_lines.append("El ancho y el alto de los bounding boxes están correlacionados. "
                                                "(Gráfica: Dispersión de Ancho vs. Alto de Bounding Boxes)")

            for i, line in enumerate(conclusion_lines, 1):
                doc.add_paragraph(line, style="List Number")

            if class_distribution:
                doc.add_paragraph()
                doc.add_paragraph("Distribución de etiquetas por clase:", style="Normal")
                for name, count in class_distribution.items():
                    p = doc.add_paragraph(f"{name}: {count} etiquetas", style="List Bullet")
                    p.paragraph_format.left_indent = Inches(0.5)
                    
            # Agregar el índice al inicio
            self.add_index(doc, section_titles)

            doc.save(self.word_path)
        
        except Exception as e:
            logger.exception(f"Error al generar el documento Word: {e}")

    def convert_to_pdf(self, word_path):
        try:
            # Intenta usar docx2pdf
            convert(word_path)
        except Exception as e:
            logger.warning(f"No se pudo convertir a PDF: {e}")
            # Para Linux, se puede configurar algo como esto:
            # try:
            #     logger.warning("Word no disponible, intentando con LibreOffice...")
            #     subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", word_path], check=True)
            # except Exception as e:
            #     logger.warning(f"No se pudo convertir a PDF: {e}")

    def run_validation(self, data_yaml):
        self.data_yaml = data_yaml
        try:
            logger.info(f"Modelo: {self.dataset_type}")
            logger.info(f"Iniciando validación...")
            logger.info("Validando distribución de clases...")
            class_distribution = self.analyze_class_distribution()
            logger.info("Analizando tamaños de imágenes...")
            image_widths, image_heights = self.analyze_image_sizes()
            logger.info("Analizando áreas de bounding boxes...")
            bbox_areas = self.analyze_bbox_areas()
            logger.info("Analizando relaciones de aspecto de imágenes...")
            aspect_ratios = self.analyze_aspect_ratios()
            logger.info("Detectando imágenes duplicadas y overlaps...")
            duplicates = self.detect_duplicates_and_overlaps()
            logger.info("Validando formato de datos con Ultralytics...")
            validation_results = self.validate_yolo_format(data_yaml)
            logger.info("Validando calidad de imágenes...")
            corrupt_images, small_images = self.validate_image_quality()
            logger.info("Analizando relaciones de aspecto de bounding boxes...")
            bbox_aspect_ratios = self.analyze_bbox_aspect_ratios()
            logger.info("Analizando posiciones centrales de bounding boxes...")
            bbox_center_positions = self.analyze_bbox_center_positions()
            logger.info("Analizando ancho vs alto de bounding boxes...")
            bbox_width_height = self.analyze_bbox_width_height()
            logger.info("Generando reporte en Word...")
            self.create_word_report(
                class_distribution,
                (image_widths, image_heights),
                bbox_areas,
                aspect_ratios,
                duplicates,
                validation_results,
                corrupt_images,
                small_images,
                bbox_aspect_ratios,
                bbox_center_positions,
                bbox_width_height,
            )
            logger.info("Generando reporte en PDF...")
            self.convert_to_pdf(self.word_path)
            logger.info("Validación completada.")
        except Exception as e:
            logger.exception(f"Error durante la validación: {e}")

    @staticmethod
    def detect_dataset_type(dataset_path):
        try:
            def is_classification(split_path):
                labels_path = os.path.join(split_path, 'labels')
                if os.path.isdir(labels_path):
                    txt_files = [f for f in os.listdir(labels_path) if f.endswith('.txt')]
                    if len(txt_files) > 0:
                        return False
                if not all(os.path.isdir(os.path.join(dataset_path, split)) for split in ['train', 'val']):
                    logger.warning("El dataset no tiene la estructura correcta")
                return True

            def is_segmentation(split_path):
                images_path = os.path.join(split_path, 'images')
                labels_path = os.path.join(split_path, 'labels')
                
                if not os.path.isdir(images_path) or not os.path.isdir(labels_path):
                    return False
                
                label_files = os.listdir(labels_path)
                if not label_files:
                    return False

                label_file = random.choice(label_files)
                if label_files:
                    label_path = os.path.join(labels_path, label_file)
                    with open(label_path, 'r') as f:
                        lines = f.readlines()
                        for line in lines:
                            parts = line.split()
                            if len(parts) <= 5:
                                return False
                return True

            def is_detection(split_path):
                images_path = os.path.join(split_path, 'images')
                labels_path = os.path.join(split_path, 'labels')
                
                if not os.path.isdir(images_path) or not os.path.isdir(labels_path):
                    return False
                
                label_files = os.listdir(labels_path)
                if not label_files:
                    return False

                label_file = random.choice(label_files)
                if label_files:
                    label_path = os.path.join(labels_path, label_file)
                    with open(label_path, 'r') as f:
                        lines = f.readlines()
                        for line in lines:
                            parts = line.split()
                            if len(parts) != 5:
                                return False
                return True

            train_path = os.path.join(dataset_path, 'train')

            if is_classification(train_path):
                return CLASSIFICATION
            elif is_segmentation(train_path):
                return SEGMENTATION
            elif is_detection(train_path):
                return DETECTION
            else:
                logger.exception("No se pudo determinar el tipo de dataset. Estructura no válida.")

        except Exception as e:
            logger.exception(f"Error durante la detección de tipo del dataset. {e}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="EDA YOLO")
    parser.add_argument("--dataset", required=True, help="Ruta al dataset")
    parser.add_argument("--project", required=True, help="Nombre del proyecto")
    parser.add_argument("--author", required=True, help="Nombre del autor")
    parser.add_argument("--content", required=True, help="Descripción del contenido")
    args = parser.parse_args()

    dataset_path = args.dataset
    project_name = args.project
    author_name = args.author
    content_description = args.content

    dataset_type = YOLODataValidator.detect_dataset_type(dataset_path)
    if dataset_type == CLASSIFICATION:
        image_folder = f"{dataset_path}/train"
        label_folder = None
        data_yaml = dataset_path
    else:  # detección y segmentación
        image_folder = f"{dataset_path}/train/images"
        label_folder = f"{dataset_path}/train/labels"
        data_yaml = f"{dataset_path}/data.yaml"

    validator = YOLODataValidator(image_folder, label_folder, data_yaml, dataset_type, dataset_path, project_name, author_name, content_description)
    validator.run_validation(data_yaml)




"""
Installing required modules:
pip install -r requirements.txt

Execution example:
py .\eda.py --dataset ".\datasets eda\clasificacion\Eyesdcar_clasificacion" --project "EyesDcar" --author "ManuC" --content "Esto es una prueba"

"""